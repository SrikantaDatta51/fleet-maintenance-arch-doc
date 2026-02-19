#!/usr/bin/env python3
"""
Healthy Fleet Maintenance & Tenant-Safe Node Lifecycle
Architecture Design Document Generator — V3

Restructured document with:
- 4 updated pillars (Phase 1 fleet ops, Phase 2 lifecycle+cert merged, NEW image pipeline, tenant)
- Combined lifecycle+NPD diagram
- Expanded SOP certification (Day 0 + Day 2)
- Better document flow
- 0.5" margins, ~15 pages
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS_DIR = os.path.join(SCRIPT_DIR, "diagrams")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "Healthy_Fleet_Maintenance_Architecture_Design.docx")

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
ACCENT_BLUE = RGBColor(0x35, 0x72, 0xA5)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def apply_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        f'  <w:start w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        f'  <w:end w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, font_size=8):
    for cell in row.cells:
        set_cell_shading(cell, "1B3A5C")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True
                r.font.size = Pt(font_size)
                r.font.name = "Calibri"


def style_data_row(row, index, font_size=7.5):
    bg = "E8F0F8" if index % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_shading(cell, bg)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            for r in p.runs:
                r.font.size = Pt(font_size)
                r.font.name = "Calibri"
                r.font.color.rgb = DARK_GRAY


def H(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    heading.paragraph_format.space_after = Pt(2)
    for r in heading.runs:
        r.font.color.rgb = DARK_BLUE
        r.font.name = "Calibri"
    return heading


def P(doc, text, bold=False, italic=False, sz=9):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(sz)
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(1)
    return para


def B(doc, text, bold_prefix="", sz=8.5):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        rb = para.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = "Calibri"
        rb.font.size = Pt(sz)
        rb.font.color.rgb = DARK_GRAY
    r = para.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(sz)
    r.font.color.rgb = DARK_GRAY
    return para


def D(doc, image_path, caption, width=Inches(5.5)):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    run = para.add_run()
    if os.path.exists(image_path):
        run.add_picture(image_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.name = "Calibri"
    cr.font.size = Pt(8)
    cr.font.color.rgb = MEDIUM_GRAY
    cr.italic = True
    cap.paragraph_format.space_after = Pt(4)


def T(doc, headers, rows, col_widths=None, hdr_sz=7.5, data_sz=7):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        pr = c.paragraphs[0]
        r = pr.add_run(hd)
        r.bold = True
        r.font.name = "Calibri"
    style_header_row(t.rows[0], hdr_sz)
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            pr = c.paragraphs[0]
            r = pr.add_run(str(v))
            r.font.name = "Calibri"
        style_data_row(t.rows[ri + 1], ri, data_sz)
    apply_table_borders(t)
    if col_widths:
        for row in t.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(1)
    return t


def main():
    print("Generating V3 Architecture Document...")
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(1)

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = DARK_BLUE
        hs.paragraph_format.space_before = Pt(6 if i == 1 else 4)
        hs.paragraph_format.space_after = Pt(2)

    # =====================================================================
    # COVER PAGE
    # =====================================================================
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Healthy Fleet Maintenance\n& Tenant-Safe Node Lifecycle")
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_BLUE
    r.font.name = "Calibri"
    r.bold = True

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Architecture Design Document")
    r.font.size = Pt(16)
    r.font.color.rgb = MEDIUM_BLUE
    r.font.name = "Calibri"

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sep.add_run("━" * 80)
    r.font.color.rgb = ACCENT_BLUE
    r.font.size = Pt(8)

    meta = [
        ("Classification", "Internal — Engineering"),
        ("Version", "3.0"),
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Author", "Distinguished Engineer, AI Compute Platform"),
        ("Status", "Final"),
        ("Review Board", "Compute Platform · K8s · OE · Storage · Network · DC Infra"),
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        ck, cv = mt.rows[i].cells[0], mt.rows[i].cells[1]
        ck.text = ""
        cv.text = ""
        pk = ck.paragraphs[0]
        pk.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.name = "Calibri"
        rk.font.size = Pt(9)
        rk.font.color.rgb = DARK_BLUE
        pv = cv.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.name = "Calibri"
        rv.font.size = Pt(9)
        rv.font.color.rgb = DARK_GRAY
        ck.width = Inches(2.0)
        cv.width = Inches(5.0)
    apply_table_borders(mt)
    doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS
    # =====================================================================
    H(doc, "Table of Contents", 1)
    for item in [
        "1.  Executive Summary",
        "2.  Pillar 1 — Healthy Fleet Maintenance (Phase 1)",
        "3.  Pillar 2 — Automated Node Lifecycle & Certification (Phase 2)",
        "4.  Pillar 3 — Automated Image Pipeline (Phase 2)",
        "5.  Pillar 4 — Tenant Reassignment",
        "6.  Node Certification SOP (\"Good Host\" Standard)",
        "7.  NPD as a Platform — Multi-Team Model",
        "8.  Node Lifecycle with NPD Detection & Enforcement",
        "9.  Scope Definition",
        "10. End-to-End Test Plan",
        "11. Deployment, Observability & Operations",
        "12. Enterprise Promises & Customer Impact",
        "Appendix A: Glossary",
    ]:
        tp = doc.add_paragraph(item)
        tp.paragraph_format.space_after = Pt(1)
        for r in tp.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            r.font.color.rgb = DARK_GRAY
    doc.add_page_break()

    # =====================================================================
    # 1. EXECUTIVE SUMMARY
    # =====================================================================
    H(doc, "1. Executive Summary", 1)
    P(doc, "This document defines the Healthy Fleet Maintenance & Tenant-Safe Node Lifecycle platform feature for "
           "the AI Compute Platform. It establishes a deterministic, automated state transition model for GPU compute "
           "nodes, ensuring every node meets a codified certification standard before serving production tenants.")

    P(doc, "The platform operates across four pillars, phased for execution:")
    B(doc, " Establish a 99.5% availability SLA squad with runbooks, observability, RMA workflows, vendor SLA "
           "tracking, and automated executive dashboards. Integrate Incident 5337 learnings. Already operational.", bold_prefix="Phase 1 — Healthy Fleet Maintenance:")
    B(doc, " Automate node lifecycle using NPD + controllers for detection and enforcement. Merge certification "
           "pipeline (BCM burn-in, DCGM L4, NCCL, HPL) into the lifecycle state machine.", bold_prefix="Phase 2 — Automated Node Lifecycle & Certification:")
    B(doc, " Deprecate BCM image cloning. Build Packer-based image pipeline with GitOps CD, cross-AZ replication, "
           "and self-service layers for Storage, K8s, and other teams.", bold_prefix="Phase 2 — Automated Image Pipeline:")
    B(doc, " Fixed state transition: Internal Tenant A → Certified → Production Tenant X → reimage + wipe + "
           "recertify. Every tenant move traverses the full certification pipeline.", bold_prefix="Tenant Reassignment:")

    P(doc, "Design is tenant-first: unhealthy nodes never serve production tenants. Self-heal is out of scope for Phase 1.")

    D(doc, os.path.join(DIAGRAMS_DIR, "01_mindmap_overview.png"),
        "Figure 1: AI Compute Platform — Four-Pillar Fleet Operations Overview", width=Inches(6.5))
    doc.add_page_break()

    # =====================================================================
    # 2. PILLAR 1 — HEALTHY FLEET MAINTENANCE (Phase 1)
    # =====================================================================
    H(doc, "2. Pillar 1 — Healthy Fleet Maintenance (Phase 1)", 1)
    P(doc, "This pillar is already operational. A dedicated 99.5% SLA squad has been established to maintain fleet "
           "health through people-driven processes, runbooks, and daily operational rigor.")

    H(doc, "2.1 SLA Definition & Critical Path", 2)
    P(doc, "The 99.5% availability target is measured at the node-level: a node is \"available\" when it passes all "
           "Good Host certification checks and is actively serving a production tenant without blocking signals.")

    T(doc, ["SLA Component", "Scope", "Critical Path"],
        [
            ["GPU Health", "All 8 GPUs healthy, ECC clean, no critical XIDs", "DCGM metrics, XID monitoring, ECC counters"],
            ["InfiniBand", "All IB ports active, NDR 400G, error counters under threshold", "ibstat, perfquery, link flap detection"],
            ["NVLink / NVSwitch", "All NVLink5 links active, zero CRC/replay errors", "nvidia-smi nvlink, nvbandwidth"],
            ["Node Readiness", "K8s node Ready, kubelet/runtime stable", "Node conditions, kubelet health"],
            ["Storage", "NVMe SMART healthy, Weka/VAST probes passing", "NVMe smart-log, storage operator"],
            ["Tenant Network", "VLAN/routing applied, node routing agent healthy", "Node routing agent, connectivity checks"],
            ["GPU Operators", "GPU operator + network operator pods running", "Operator pod status, CRD health"],
        ],
        col_widths=[1.3, 2.5, 3.2]
    )

    H(doc, "2.2 Operational Squad & Processes", 2)
    T(doc, ["Item", "Details"],
        [
            ["Squad", "Dedicated L1 99.5% SLA squad — established and operational"],
            ["Runbooks", "Runbooks for every critical path component: GPU, IB, NVLink, OS, disk, K8s node, operator"],
            ["Daily Triage", "Daily fleet health triage with status review, owner assignment, ETA tracking"],
            ["Observability", "Grafana dashboards per critical path component; alerting pipeline operational"],
            ["Executive Dashboards", "Automated daily executive summary / business dashboards emitted"],
            ["GPU Reliability", "Q4 GPU reliability tasks completed (XID patterns, ECC trends, RMA rates)"],
            ["RMA Workflow", "Evidence capture → ticket routing → vendor SLA tracking → post-RMA recertification gate"],
            ["Vendor SLAs", "NVIDIA RMA SLAs documented and tracked; DC Infra swap SLAs committed"],
            ["Capacity Return", "Clean SOP: ensure failed node capacity returned to fleet in under 24 hours"],
        ],
        col_widths=[1.8, 5.2]
    )

    H(doc, "2.3 Incident 5337 Integration", 2)
    P(doc, "Multi-node distributed training failure learnings are integrated into Phase 1 operations:")
    B(doc, "Stricter IB readiness thresholds for multi-node workloads; link flap detection tightened")
    B(doc, "Correlation across node sets: job failure → suspect node identification → quarantine repeat offenders")
    B(doc, "Certification includes multi-node smoke checks (NCCL cross-node + sample distributed job)")
    B(doc, "Bus bandwidth AND algorithm bandwidth must both scale — validated via nccl-tests and HPL multi-node")
    B(doc, "Repeat offender tracking: nodes with >2 quarantine events in 7 days flagged for RMA triage")
    doc.add_page_break()

    # =====================================================================
    # 3. PILLAR 2 — AUTOMATED NODE LIFECYCLE & CERTIFICATION (Phase 2)
    # =====================================================================
    H(doc, "3. Pillar 2 — Automated Node Lifecycle & Certification (Phase 2)", 1)
    P(doc, "This pillar merges automated node lifecycle management with the node certification pipeline. "
           "NPD and custom controllers detect host-level degradation and enforce tenant safety. The certification "
           "pipeline (BCM burn-in, DCGM L4, NCCL, HPL) is integrated into the lifecycle state machine.")

    H(doc, "3.1 Layered Component View (Our Ownership Boundary)", 2)
    P(doc, "The Compute Platform team owns through the Kubernetes boundary. K8s, Storage, and RMA are coordination dependencies.")
    T(doc, ["Layer", "Components", "Ownership"],
        [
            ["Hardware", "GPU (B200), HCA/NIC (IB NDR 400G), NVMe, BMC, NVLink5/NVSwitch, CPU/Memory", "Compute Platform"],
            ["Host OS", "Kernel, drivers, filesystem, time sync, container runtime", "Compute Platform"],
            ["NVIDIA Stack", "BCM, DCGM/DCGMI, gpu_burn, nccl-tests, nvbandwidth, HPL", "Compute Platform"],
            ["GPU / Network Operators", "GPU operator, network operator, IB SR-IOV / VFs, Multus", "Compute Platform"],
            ["Tenant Network", "Node routing agent, VLAN/tenant routing, IB fabric integration", "Compute Platform + Network"],
            ["Kubernetes", "Kubelet, CNI (Tigera/Calico), node readiness, scheduling", "K8s Team (coordinate)"],
            ["Storage", "Weka/VAST probes, NVMe health, storage operator", "Storage Team (coordinate)"],
            ["RMA / DC Infra", "Physical swap, shipping, vendor coordination", "DC Infra (coordinate)"],
        ],
        col_widths=[1.5, 3.3, 2.2]
    )

    H(doc, "3.2 Detection & Enforcement Model", 2)
    P(doc, "NPD detectors feed host signals into a standardized framework. On any blocking signal, the node is "
           "quarantined (cordon + taint NoSchedule). Self-heal is explicitly out of scope for Phase 1.")

    D(doc, os.path.join(DIAGRAMS_DIR, "03_workflow_diagram.png"),
        "Figure 2: Detection → Enforcement → Recovery Workflow (Self-Heal: Out of Scope)", width=Inches(6.0))

    H(doc, "3.3 Lifecycle State Machine", 2)
    P(doc, "The lifecycle has 8 states with deterministic transitions. The certification pipeline (burn-in → DCGM L4 "
           "→ NCCL → HPL) is embedded in the state machine. Every node must pass through INTERNAL_TENANT_A "
           "for certification before entering production.")

    D(doc, os.path.join(DIAGRAMS_DIR, "06_lifecycle_combined.png"),
        "Figure 3: Node Lifecycle with NPD Detection & Enforcement (8 States)", width=Inches(6.5))

    T(doc, ["State", "Description", "Entry", "Exit"],
        [
            ["INTERNAL_TENANT_A", "Certification testing environment", "New node / detach+wipe / certified", "All tests pass → CERTIFIED"],
            ["CERTIFIED", "All certification layers passed", "Recertification pass", "Admit to fleet → BUFFER_HEALTHY"],
            ["BUFFER_HEALTHY", "Certified, available for assignment", "Certified + admitted", "Tenant attach OR blocking signal"],
            ["TENANT_ASSIGNED", "Serving production tenant", "Attach tenant label", "Detach+wipe OR blocking signal"],
            ["QUARANTINED", "Blocked; removed from scheduling", "Blocking signal from any state", "→ REPAIR or → RMA"],
            ["REPAIR", "Software/config fix underway", "Triage: SW/config issue", "Fix complete → BURN_IN"],
            ["RMA", "Hardware replacement in progress", "Hardware fault confirmed", "Replaced → BURN_IN"],
            ["BURN_IN", "BCM burn-in (6–24h)", "Post-repair or post-RMA", "Pass → RECERTIFY; Fail → QUARANTINED"],
        ],
        col_widths=[1.5, 2.0, 1.8, 1.7]
    )
    doc.add_page_break()

    # =====================================================================
    # 4. PILLAR 3 — AUTOMATED IMAGE PIPELINE (Phase 2)
    # =====================================================================
    H(doc, "4. Pillar 3 — Automated Image Pipeline (Phase 2)", 1)
    P(doc, "BCM's current image cloning model does not support reproducible, cross-environment image replication. "
           "This pillar replaces it with a Packer-based, GitOps-driven image pipeline.")

    H(doc, "4.1 Current State & Problem", 2)
    B(doc, "BCM uses image cloning (golden image → clone to nodes) with no version control or audit trail")
    B(doc, "No mechanism to replicate images from one AZ/region to another consistently")
    B(doc, "Other teams (Storage, K8s) cannot self-serve their layer configuration into the base image")
    B(doc, "Re-imaging after tenant transitions requires manual intervention")

    D(doc, os.path.join(DIAGRAMS_DIR, "07_image_pipeline.png"),
        "Figure 5: Automated Image Pipeline — Packer-Based GitOps Model", width=Inches(6.5))

    H(doc, "4.2 Target Architecture", 2)
    T(doc, ["Component", "Details"],
        [
            ["Packer Pipeline", "Packer templates with NVIDIA support; builds DGX OS base image with all drivers, DCGM, NCCL, operators"],
            ["GitOps CD", "Git repository as source of truth; CI/CD pipeline triggers image builds on merge; tagged releases"],
            ["Cross-AZ Replication", "Built images replicated from primary AZ to all target AZs/regions automatically"],
            ["Self-Service Layers", "Storage team adds Weka/VAST agents; K8s team adds kubelet/CNI config; each layer has tests"],
            ["Image Versioning", "Semantic versioning; rollback support; image catalog with metadata"],
            ["BCM Integration", "BCM provisions nodes using Packer-built images; image ID tracked in BCM node record"],
        ],
        col_widths=[1.8, 5.2]
    )

    H(doc, "4.3 Layer Contributions", 2)
    T(doc, ["Layer", "Owner", "Contributes to Image"],
        [
            ["Base OS + Drivers", "Compute Platform", "Kernel, NVIDIA drivers, CUDA, container runtime, BCM agent"],
            ["DCGM + Diagnostics", "Compute Platform", "DCGM, DCGMI, gpu_burn, nccl-tests, nvbandwidth"],
            ["GPU / Network Operators", "Compute Platform", "Operator manifests, IB SR-IOV config, Multus"],
            ["Storage Agents", "Storage Team", "Weka client, VAST agent, mount configuration"],
            ["K8s Components", "K8s Team", "Kubelet config, CNI (Tigera/Calico), node labels/annotations"],
            ["Monitoring", "OE Team", "Log/metric agents, NPD config, alert rules"],
        ],
        col_widths=[1.5, 1.5, 4.0]
    )
    doc.add_page_break()

    # =====================================================================
    # 5. PILLAR 4 — TENANT REASSIGNMENT
    # =====================================================================
    H(doc, "5. Pillar 4 — Tenant Reassignment", 1)
    P(doc, "Every node movement between tenants follows a fixed, secure state transition. There are no spare "
           "buffer nodes — all nodes are assigned to Internal Tenant A for continuous certification testing when "
           "not serving production tenants.")

    H(doc, "5.1 Reassignment Lifecycle", 2)
    D(doc, os.path.join(DIAGRAMS_DIR, "05_buffer_strategy.png"),
        "Figure 4: Internal Tenant Certification & Fleet Lifecycle Model", width=Inches(6.0))

    P(doc, "The fixed transition is: Tenant X → detach → reimage + disk wipe → Internal Tenant A (certification) → "
           "Certified → Tenant Y. No shortcuts are permitted. This ensures complete tenant isolation and fresh certification.")

    H(doc, "5.2 Reassignment Steps", 2)
    T(doc, ["Step", "Action", "Validation"],
        [
            ["1. Detach", "Remove tenant label, drain workloads", "No tenant pods running on node"],
            ["2. Reimage", "Full OS reimage from Packer-built golden image", "Image hash matches catalog; BCM confirms"],
            ["3. Disk Wipe", "Cryptographic wipe of all local storage", "Wipe verification log; SMART health OK"],
            ["4. Network Reconfig", "Revert tenant VLAN/routing; apply base network config", "IB ports active; base routing validated"],
            ["5. Certification", "Full Day Zero SOP in Internal Tenant A", "All L1–L9 certification layers pass"],
            ["6. Assign", "Attach new tenant label; apply tenant network config", "Tenant routing validated; scheduling enabled"],
        ],
        col_widths=[0.8, 2.8, 3.4]
    )

    H(doc, "5.3 Internal Tenant A — Daily Certification", 2)
    P(doc, "Nodes in Internal Tenant A run daily certification tests including multi-node sample jobs. "
           "This serves as both continuous validation and the certification gate for tenant assignment. "
           "Tests validate bus bandwidth and algorithm bandwidth scaling per Incident 5337 retrospective.")
    B(doc, "Daily burn-in cycles with gpu_burn at various precisions (FP64/FP32/TF32/FP8)")
    B(doc, "NCCL all_reduce_perf across node sets — bus BW must meet ≥1530 GB/s (85% of NVLink5 1.8 TB/s)")
    B(doc, "Multi-node sample job: validates cross-node IB communication, NCCL collective scaling")
    B(doc, "K8s dummy job: validates pod scheduling, GPU device plugin, operator health")
    doc.add_page_break()

    # =====================================================================
    # 6. NODE CERTIFICATION SOP ("Good Host" Standard)
    # =====================================================================
    H(doc, "6. Node Certification SOP (\"Good Host\" Standard)", 1)
    P(doc, "Aligned with SOP-GPU-CERT-001 v1.1 for NVIDIA B200 GPU nodes. A node must pass ALL certification "
           "layers before being flagged CERTIFIED and eligible for production tenant assignment.")

    H(doc, "6.1 Day Zero — Certification Layers", 2)

    T(doc, ["Layer", "Test", "Tool / Command", "B200 Pass Criteria"],
        [
            ["L1: Pre-Flight", "GPU discovery (8× B200)", "nvidia-smi -L", "8 GPUs detected"],
            ["L1: Pre-Flight", "Driver / CUDA version", "nvidia-smi", "Matches BCM baseline"],
            ["L1: Pre-Flight", "DCGM agent running", "systemctl status nvidia-dcgm", "Active/Running"],
            ["L1: Pre-Flight", "NVSwitch discovery", "dcgmi discovery -l", "All NVSwitches detected"],
            ["L1: Pre-Flight", "ECC enabled (all GPUs)", "nvidia-smi -q -d ECC", "ON — all GPUs"],
            ["L1: Pre-Flight", "Firmware versions", "nvfwupd show_version", "Matches fleet baseline"],
            ["L1: Pre-Flight", "System health", "sudo nvsm show health", "Healthy"],
            ["L1: Pre-Flight", "InfiniBand ports (NDR 400G)", "ibstat", "All ports Active"],
            ["L1: Pre-Flight", "PCIe link width/speed", "lspci -vvv", "Gen5 x16, no downgrade"],
            ["L2: GPU Burn-In", "Sustained GPU stress (100% util)", "gpu_burn -tc 21600", "Zero XID, zero DBE, 6–24h"],
            ["L2: GPU Burn-In", "Temperature monitoring", "nvidia-smi dmon", "≤70°C sustained (liquid-cooled)"],
            ["L2: GPU Burn-In", "Power monitoring", "nvidia-smi dmon", "≤1000W TDP per GPU"],
            ["L3: DCGM L4", "All plugins: HW, SM stress, memtest, EUD", "dcgmi diag -r 4", "ALL plugins PASS (~90 min)"],
            ["L3: DCGM L4", "NVLink5 bandwidth", "Integration plugin", "BW ≥90% of 1.8 TB/s"],
            ["L3: DCGM L4", "HBM3e memtest (192 GB/GPU)", "Memtest plugin", "Zero memory errors"],
            ["L4: NCCL", "All-Reduce (intra-node 8 GPUs)", "all_reduce_perf -g 8", "BusBW ≥1530 GB/s"],
            ["L4: NCCL", "All-Gather, Reduce-Scatter, Broadcast", "nccl-tests suite", "BusBW ≥1530 GB/s"],
            ["L4: NCCL", "Multi-node cross-node all-reduce", "mpirun + all_reduce_perf", "≥85% per-port IB line rate"],
            ["L5: HPL", "FP64 peak FLOPS", "HPL (NGC container)", "≥70% theoretical peak"],
            ["L5: HPL", "Mixed-precision (HPL-MxP)", "HPL-MxP (NGC)", "≥70% theoretical peak"],
            ["L6: NVLink/NVSwitch", "NVLink status + errors", "nvidia-smi nvlink --status/-e", "All active, zero CRC/replay"],
            ["L6: NVLink/NVSwitch", "Device-to-device bandwidth", "nvbandwidth", "D2D ≥90% of 1.8 TB/s"],
            ["L6: IB Network", "IB port status", "ibstat", "Active, NDR 400 Gb/s"],
            ["L6: IB Network", "RDMA write bandwidth", "ib_write_bw", "≥90% line rate"],
            ["L6: IB Network", "GPUDirect RDMA", "Multi-node NCCL", "≥85% per-port line rate"],
            ["L7: Storage", "NVMe SMART health", "nvme smart-log /dev/nvme*", "Spare ≥80%, zero media errors"],
            ["L7: Storage", "NVMe throughput", "fio (seq + random)", "Meets drive spec"],
            ["L7: System", "System memory", "memtester", "Zero errors"],
            ["L7: System", "CPU stress", "stress-ng --cpu $(nproc)", "Zero failures, no MCEs"],
            ["L7: System", "BMC sensors + PSU", "ipmitool sensor/sdr list", "All in range, all OK"],
            ["L8: Multi-Node Job", "Sample distributed training job", "mpirun + training script", "Completes, no errors"],
            ["L8: Multi-Node Job", "Cross-node connectivity", "IB health + NCCL scaling", "Bus BW + algo BW scale"],
            ["L9: K8s Dummy Job", "Pod scheduling + GPU allocation", "kubectl + sample workload", "Job completes, GPUs visible"],
        ],
        col_widths=[1.2, 1.8, 2.0, 2.0]
    )

    P(doc, "Day Zero Gate: Node admitted ONLY if ALL layers pass. Output: certification record → node marked CERTIFIED.", bold=True, sz=8.5)

    H(doc, "6.2 Day Two — Continuous Monitoring", 2)
    T(doc, ["Metric", "Source", "B200 Healthy", "Severity"],
        [
            ["GPU Temp", "DCGM_FI_DEV_GPU_TEMP", "≤70°C (liquid)", "P2 if exceeded"],
            ["HBM Temp", "DCGM_FI_DEV_MEMORY_TEMP", "≤95°C", "P2 if exceeded"],
            ["Power Draw", "DCGM_FI_DEV_POWER_USAGE", "≤1000W", "P2 if exceeded"],
            ["ECC SBE", "DCGM_FI_DEV_ECC_SBE_VOL_TOTAL", "0", "P2 — investigate 4h"],
            ["ECC DBE", "DCGM_FI_DEV_ECC_DBE_VOL_TOTAL", "Must be 0", "P1 — immediate quarantine"],
            ["Retired Pages (DBE)", "DCGM_FI_DEV_RETIRED_DBE", "Must be 0", "P1 — immediate quarantine"],
            ["Row Remap Failure", "DCGM_FI_DEV_ROW_REMAP_FAILURE", "Must be 0", "P1 — immediate quarantine"],
            ["XID Errors", "DCGM_FI_DEV_XID_ERRORS", "0", "P1 for critical XIDs (48/63/64/79/94/95)"],
            ["NVLink CRC", "DCGM_FI_DEV_NVLINK_CRC_*", "0", "P2 — investigate"],
            ["PCIe Replay", "DCGM_FI_DEV_PCIE_REPLAY_COUNTER", "0 per 24h", "P3 — weekly review"],
            ["IB Errors", "perfquery", "0 per 24h", "P2 — investigate"],
            ["NVMe Spare", "smartctl", "≥50%", "P3 if <80%; P2 if <50%"],
        ],
        col_widths=[1.3, 2.2, 1.5, 2.0]
    )

    H(doc, "6.3 Periodic Re-Validation", 2)
    T(doc, ["Frequency", "Test", "Duration"],
        [
            ["Weekly", "dcgmi diag -r 1 (quick health)", "~2 min"],
            ["Monthly", "dcgmi diag -r 3 (medium diagnostic)", "~15 min"],
            ["Quarterly", "dcgmi diag -r 4 + NCCL collectives + NVLink", "~100 min"],
            ["Annually", "Full burn-in + HPL + 24h soak (same as Day Zero)", "~8 hrs"],
        ],
        col_widths=[1.5, 3.5, 2.0]
    )
    doc.add_page_break()

    # =====================================================================
    # 7. NPD AS A PLATFORM
    # =====================================================================
    H(doc, "7. NPD as a Platform — Multi-Team Contribution Model", 1)
    P(doc, "NPD is the common substrate for node health detection. Five teams contribute domain-specific detectors "
           "under a single framework with standardized outputs.")

    D(doc, os.path.join(DIAGRAMS_DIR, "04_npd_platform.png"),
        "Figure 6: NPD as a Platform — Multi-Team Contribution Model", width=Inches(5.5))

    T(doc, ["Team", "Owns / Contributes to NPD", "Specific Ask"],
        [
            ["Compute Platform", "Host-side IB & tenant VLAN; GPU/Network Operators; IB SR-IOV/VFs; DCGM integration; GPU health & XID detection; host-level checks",
             "Owns NPD framework, detector development, certification pipeline, state machine controller"],
            ["Kubernetes", "Kubelet/runtime; node readiness; CNI (Tigera/Calico); nodefs/imagefs pressure",
             "CNI health probes, node readiness signal fidelity, taint/toleration policy enforcement"],
            ["Storage", "Weka probes; VAST probes",
             "Storage health probe integration into NPD; R/W latency signals; capacity alerts"],
            ["OE", "NPD integration support; logs, events, metrics across all AZs",
             "Log/metric pipeline reliability; NPD deployment across all AZs; monitoring infra"],
            ["Network / App Network", "Ingress probes for inference workloads",
             "Ingress health signals; endpoint readiness probes; latency thresholds for inference"],
        ],
        col_widths=[1.3, 2.8, 2.9]
    )

    # =====================================================================
    # 8. SCOPE DEFINITION (was section 9, moved up for flow)
    # =====================================================================
    H(doc, "8. Scope Definition", 1)
    T(doc, ["Category", "In Scope", "Out of Scope"],
        [
            ["People / Process", "L1 squad, runbooks, daily triage, escalation, exec dashboards", "Tenant application ops"],
            ["Host Certification", "BCM burn-in, DCGM L4, NCCL, HPL, ECC, NVLink, IB, storage checks", "Tenant workload tests beyond host gating"],
            ["Node Lifecycle", "8-state machine, cordon/taint, deterministic reassignment", "Auto-reboot/reimage without gate"],
            ["Self-Heal", "Out of scope for Phase 1", "Automated self-remediation (future)"],
            ["Image Pipeline", "Packer build, GitOps CD, cross-AZ replication, self-serve layers", "BCM image cloning (deprecated)"],
            ["Our Boundary", "Hardware through GPU/Network operators and tenant networking", "K8s scheduling policy, storage backend ops"],
            ["K8s (Coordinate)", "Node readiness signals, CNI health, taint enforcement", "K8s control plane, scheduler logic"],
            ["Storage (Coordinate)", "Weka/VAST probes, NVMe health integration", "Storage backend operations"],
            ["RMA (Coordinate)", "Evidence capture, ticket routing, recertification gate", "Physical swap (DC Infra + vendor)"],
        ],
        col_widths=[1.5, 2.7, 2.8]
    )

    # =====================================================================
    # 9. TEST PLAN (compact)
    # =====================================================================
    H(doc, "9. End-to-End Test Plan", 1)
    T(doc, ["Test Layer", "What We Validate", "Where"],
        [
            ["Certification (D0)", "All L1–L9 SOP layers pass", "Internal Tenant A"],
            ["Continuous (D2)", "DCGM metrics, IB health, NVMe, operators healthy", "All nodes, always"],
            ["Multi-Node", "NCCL cross-node, bus BW + algo BW scaling, sample job", "Internal Tenant A node sets"],
            ["Enforcement", "Taint/cordon on blocking signal; scheduling blocked", "Continuous (NPD)"],
            ["Recovery", "Repair → burn-in → recertify → CERTIFIED → fleet", "Post-fix pipeline"],
            ["Reassignment", "Reimage + wipe + full recertification + network reconfig", "Every tenant transition"],
            ["Image Pipeline", "Packer build reproducible; image replicates across AZs", "CI/CD pipeline"],
        ],
        col_widths=[1.5, 3.5, 2.0]
    )

    H(doc, "Acceptance Criteria", 2)
    B(doc, " Node cannot enter TENANT_ASSIGNED unless all L1–L9 certification layers pass.", bold_prefix="AC-1:")
    B(doc, " Blocking signal → quarantine (cordon+taint) within detection-to-quarantine SLO.", bold_prefix="AC-2:")
    B(doc, " No return-to-fleet without full recertification (burn-in + DCGM L4 + NCCL + HPL).", bold_prefix="AC-3:")
    B(doc, " Tenant reassignment always follows: reimage → disk wipe → recertification → network reconfig.", bold_prefix="AC-4:")
    B(doc, " Bus BW ≥1530 GB/s and algo BW must scale for distributed workload eligibility.", bold_prefix="AC-5:")
    B(doc, " Packer-built images replicate to all target AZs before any deployment.", bold_prefix="AC-6:")

    # =====================================================================
    # 10. DEPLOYMENT & OBSERVABILITY (compact)
    # =====================================================================
    H(doc, "10. Deployment, Observability & Operations", 1)
    H(doc, "10.1 Staged Rollout", 2)
    T(doc, ["Stage", "Action", "Risk"],
        [
            ["1", "Deploy NPD baseline + dashboards (observe only)", "No enforcement"],
            ["2", "Enable enforcement in Internal Tenant A only", "Limited blast radius"],
            ["3", "Enforce certification gating for production tenant attachment", "Tenant-impacting; rollback plan"],
            ["4", "Enable automated reassignment + image pipeline", "Full automation; max monitoring"],
        ],
        col_widths=[0.5, 4.0, 2.5]
    )

    H(doc, "10.2 Dashboards & Alerts", 2)
    P(doc, "Dashboards:", bold=True, sz=8.5)
    B(doc, "Fleet by state: Internal Tenant A, CERTIFIED, BUFFER_HEALTHY, TENANT_ASSIGNED, QUARANTINED, BURN_IN, RMA")
    B(doc, "Detection-to-quarantine latency (P50/P95/P99) · Quarantine-to-recovery latency · Certification pass/fail rate")
    B(doc, "NCCL/HPL benchmark trends · Repeat offender tracking · Daily automated executive summary")
    P(doc, "Alerts:", bold=True, sz=8.5)
    B(doc, "Blocking taint/cordon applied · Certification failure · Multi-node failure correlation spike")
    B(doc, "Vendor SLA breach risk · Image pipeline build failure · Cross-AZ replication failure")

    H(doc, "10.3 L1 Daily Report", 2)
    T(doc, ["Item", "Content"],
        [
            ["Unhealthy Nodes", "Hostname, fault category, owner, ETA"],
            ["Returned to Fleet", "Nodes returned in last 24h, cause summary"],
            ["Escalated (RMA)", "Evidence status, vendor ticket #, SLA tracking"],
            ["Certification Pipeline", "Throughput, pass/fail rate, average time-to-certify"],
            ["Fleet Health", "Total fleet, % TENANT_ASSIGNED, % in certification, % quarantined"],
        ],
        col_widths=[1.8, 5.2]
    )

    # =====================================================================
    # 11. ENTERPRISE PROMISES & CUSTOMER IMPACT (combined, compact)
    # =====================================================================
    H(doc, "11. Enterprise Promises & Customer Impact", 1)
    T(doc, ["#", "Promise", "Delivered By"],
        [
            ["1", "Tenant Safety by Default", "NPD + taint enforcement; unhealthy nodes never serve production tenants"],
            ["2", "Predictable Availability (99.5%)", "Dedicated SLA squad, runbooks, <24h capacity return, automated exec dashboards"],
            ["3", "Deterministic Lifecycle", "8-state machine; all transitions follow secure, clean, auditable path"],
            ["4", "Certified Hosts Only", "L1–L9 certification SOP; DCGM L4, NCCL bus/algo BW, HPL benchmarks"],
            ["5", "Reproducible Images", "Packer pipeline, GitOps CD, cross-AZ replication, self-serve layers"],
            ["6", "Platform Extensibility", "NPD as shared platform; 5 teams contribute detectors independently"],
        ],
        col_widths=[0.4, 2.0, 4.6]
    )

    H(doc, "Customer Impact", 2)
    P(doc, "Positive:", bold=True, sz=8.5)
    B(doc, "Fewer tenant-visible failures; higher multi-node job success rate (5337 mitigation)")
    B(doc, "Cleaner tenant isolation (reimage + wipe); consistent image across AZs")
    B(doc, "Layer-by-layer certification ensures bus BW and algo BW scale before production exposure")
    P(doc, "Tradeoffs:", bold=True, sz=8.5)
    B(doc, "Stricter gating may temporarily reduce capacity; quarantine-first (no self-heal) increases repair dependency in Phase 1")
    doc.add_page_break()

    # =====================================================================
    # APPENDIX: GLOSSARY
    # =====================================================================
    H(doc, "Appendix A: Glossary", 1)
    T(doc, ["Term", "Definition"],
        [
            ["BCM", "NVIDIA Base Command Manager — cluster management and bare-metal provisioning"],
            ["DCGM / DCGMI", "NVIDIA Data Center GPU Manager / CLI Interface — GPU diagnostics and monitoring"],
            ["ECC / DBE / SBE", "Error Correcting Code / Double-Bit Error / Single-Bit Error — memory error categories"],
            ["gpu_burn", "GPU stress testing tool — sustained matrix multiply at 100% utilization"],
            ["HPL / HPL-MxP", "High Performance Linpack — FP64 and mixed-precision benchmark for GPU compute"],
            ["IB / NDR", "InfiniBand / Next Data Rate (400 Gb/s) — high-performance networking fabric"],
            ["NCCL", "NVIDIA Collective Communications Library — GPU-to-GPU collective operations"],
            ["NPD", "Node Problem Detector — Kubernetes DaemonSet for node health monitoring"],
            ["NVLink5", "NVIDIA GPU interconnect — 1.8 TB/s per GPU bidirectional (B200)"],
            ["Packer", "HashiCorp tool for building automated machine images from configuration"],
            ["RMA", "Return Merchandise Authorization — hardware replacement process with vendor"],
            ["SR-IOV / VF", "Single Root I/O Virtualization / Virtual Function — hardware network virtualization"],
            ["XID", "NVIDIA GPU error identifier code — categorizes GPU fault types"],
        ],
        col_widths=[1.5, 5.5]
    )

    doc.save(OUTPUT_FILE)
    fsize = os.path.getsize(OUTPUT_FILE)
    print(f"\n✅ Saved: {OUTPUT_FILE}")
    print(f"   Size: {fsize/1024:.1f} KB")


if __name__ == "__main__":
    main()
